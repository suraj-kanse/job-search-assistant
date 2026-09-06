import os
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import zipfile
import shutil
from profile_manager import load_profile

def tailor_profile_data(job, profile):
    """Tailors skills, projects, and experiences dynamically based on the job description/snippet."""
    import re
    
    jd_text = (job["title"] + " " + job.get("snippet", "")).lower()
    
    # Deep copy profile data so we don't mutate the original loaded JSON
    tailored = json.loads(json.dumps(profile))
    
    # 1. Re-order Projects by keyword overlap
    projects = tailored.get("projects", [])
    if projects:
        for proj in projects:
            score = 0
            proj_text = (proj["name"] + " " + proj["tech_stack"] + " " + " ".join(proj["details"])).lower()
            for word in re.findall(r'\w+', jd_text):
                if len(word) > 2 and word in proj_text:
                    score += 1
            proj["_score"] = score
            
            # Re-order the bullet points inside the project
            details = proj["details"]
            def score_bullet(b):
                b_score = 0
                for w in re.findall(r'\w+', jd_text):
                    if len(w) > 2 and w in b.lower():
                        b_score += 1
                return b_score
            proj["details"] = sorted(details, key=score_bullet, reverse=True)
            
        tailored["projects"] = sorted(projects, key=lambda p: p["_score"], reverse=True)

    # 2. Re-order Experience bullet points
    experience = tailored.get("experience", [])
    if experience:
        for exp in experience:
            details = exp["details"]
            def score_bullet(b):
                b_score = 0
                for w in re.findall(r'\w+', jd_text):
                    if len(w) > 2 and w in b.lower():
                        b_score += 1
                return b_score
            exp["details"] = sorted(details, key=score_bullet, reverse=True)

    # 3. Re-order and highlight skills
    skills = tailored.get("skills", {})
    tailored_skills = {}
    cat_scores = {}
    
    for cat, list_skills in skills.items():
        score = 0
        skill_scores = []
        for s in list_skills:
            s_score = 0
            if s.lower() in jd_text:
                s_score = 5
            elif any(w in s.lower() for w in re.findall(r'\w+', jd_text) if len(w) > 2):
                s_score = 2
            skill_scores.append((s, s_score))
            score += s_score
            
        cat_scores[cat] = score
        tailored_skills[cat] = [item[0] for item in sorted(skill_scores, key=lambda x: x[1], reverse=True)]
        
    sorted_cats = sorted(cat_scores.keys(), key=lambda c: cat_scores[c], reverse=True)
    rebuilt_skills = {}
    for cat in sorted_cats:
        rebuilt_skills[cat] = tailored_skills[cat]
        
    tailored["skills"] = rebuilt_skills
    return tailored

def select_relevant_items(items, jd_text, min_k=2, max_k=4):
    """Ranks and selects the top 2-4 most impactful items based on JD keyword relevance."""
    if not items:
        return []
    import re
    words = [w.lower() for w in re.findall(r'\w+', jd_text) if len(w) > 2]
    
    scored = []
    for item in items:
        score = 0
        item_lower = item.lower()
        for w in words:
            if w in item_lower:
                score += 1
        scored.append((item, score))
        
    ranked = [item for item, s in sorted(scored, key=lambda x: x[1], reverse=True)]
    return ranked[:max(min_k, min(max_k, len(ranked)))]

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, url, text, color="1D63B8", underline=False):
    """Adds a native clickable hyperlink to a docx paragraph that works in Word and PDF exports."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def create_resume(job, profile, output_path):
    """Generates an ATS-compliant resume calibrated strictly for 1-page PDF export with clickable links."""
    doc = docx.Document()
    jd_text = (job["title"] + " " + job.get("snippet", "")).lower()
    
    # Page setup - 0.48 inch compact margins to guarantee strict 1-page PDF export
    for section in doc.sections:
        section.top_margin = Inches(0.48)
        section.bottom_margin = Inches(0.48)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

    # Base styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(9.2)

    # 1. Header (Centered Name & Contact with clickable links)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(1)
    run = title_p.add_run(profile["name"].upper())
    run.bold = True
    run.font.size = Pt(17)
    
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_before = Pt(0)
    contact_p.paragraph_format.space_after = Pt(4)
    
    contact_p.add_run(f"{profile['phone']} - ")
    add_hyperlink(contact_p, f"mailto:{profile['email']}", profile['email'])
    contact_p.add_run(f" - {profile['location']['current']}\nLinkedIn: ")
    add_hyperlink(contact_p, f"https://{profile['linkedin']}", profile['linkedin'])
    contact_p.add_run(" - GitHub: ")
    add_hyperlink(contact_p, f"https://{profile['github']}", profile['github'])

    # Helper function for standardized section headings
    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4.5)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10.5)
        
    # 2. Professional Summary
    add_section_heading("Professional Summary")
    summary_p = doc.add_paragraph()
    summary_p.paragraph_format.space_before = Pt(0)
    summary_p.paragraph_format.space_after = Pt(2)
    summary_text = (
        "Information Technology undergraduate with hands-on experience building full-stack and "
        "backend applications using Python, Django, JavaScript, React, Node.js, and SQL/NoSQL databases. "
        "Developed and deployed web applications featuring REST APIs, authentication, role-based access control, "
        "database integration, and automated data/resume processing through internship and academic projects."
    )
    summary_p.add_run(summary_text)

    # 3. Technical Skills
    add_section_heading("Technical Skills")
    skills = profile.get("skills", {})
    categories = [
        ("Languages", skills.get("languages", [])),
        ("Backend", skills.get("backend", [])),
        ("Frontend", skills.get("frontend", [])),
        ("Databases", skills.get("databases", [])),
        ("Security", skills.get("security", [])),
        ("Tools", skills.get("tools", []))
    ]
    for cat_name, list_skills in categories:
        if list_skills:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.1)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0.5)
            run_cat = p.add_run(f"{cat_name}: ")
            run_cat.bold = True
            p.add_run(", ".join(list_skills))

    # 4. Internship Experience
    add_section_heading("Internship Experience")
    for exp in profile.get("experience", []):
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(1.5)
        p_title.paragraph_format.space_after = Pt(0.5)
        run_role = p_title.add_run(f"{exp['role']} - {exp['company']} ({exp['duration']})")
        run_role.bold = True
        
        for detail in exp.get("details", []):
            p_bullet = doc.add_paragraph(style='List Bullet')
            p_bullet.paragraph_format.space_before = Pt(0)
            p_bullet.paragraph_format.space_after = Pt(0.5)
            p_bullet.paragraph_format.left_indent = Inches(0.2)
            p_bullet.add_run(detail)

    # 5. Projects
    add_section_heading("Projects")
    for proj in profile.get("projects", []):
        p_proj = doc.add_paragraph()
        p_proj.paragraph_format.space_before = Pt(2)
        p_proj.paragraph_format.space_after = Pt(0.5)
        run_proj = p_proj.add_run(f"{proj['name']} ({proj['duration']}) - Tech Stack: {proj['tech_stack']}")
        run_proj.bold = True
        
        for detail in proj.get("details", []):
            p_bullet = doc.add_paragraph(style='List Bullet')
            p_bullet.paragraph_format.space_before = Pt(0)
            p_bullet.paragraph_format.space_after = Pt(0.5)
            p_bullet.paragraph_format.left_indent = Inches(0.2)
            p_bullet.add_run(detail)

        # Include clickable project links if present
        links = proj.get("links", {})
        if links:
            p_link = doc.add_paragraph()
            p_link.paragraph_format.space_before = Pt(0)
            p_link.paragraph_format.space_after = Pt(1)
            p_link.paragraph_format.left_indent = Inches(0.2)
            first = True
            for k, v in links.items():
                if not first:
                    p_link.add_run("  -  ")
                p_link.add_run(f"{k.capitalize()}: ")
                if v.startswith("http"):
                    add_hyperlink(p_link, v, v, color="1D63B8", underline=True)
                else:
                    p_link.add_run(v)
                first = False

    # 6. Education
    add_section_heading("Education")
    for edu in profile.get("education", []):
        p_edu = doc.add_paragraph()
        p_edu.paragraph_format.space_before = Pt(0)
        p_edu.paragraph_format.space_after = Pt(1.5)
        run_deg = p_edu.add_run(f"{edu['degree']} - {edu['college']} | {edu['duration']} | CGPA: {edu['cgpa']}")
        run_deg.bold = True

    # 7. Certifications (Smart selection of top 2-4 relevant)
    certs = profile.get("certifications", [])
    if certs:
        selected_certs = select_relevant_items(certs, jd_text, min_k=2, max_k=4)
        add_section_heading("Certifications")
        p_cert = doc.add_paragraph()
        p_cert.paragraph_format.left_indent = Inches(0.1)
        p_cert.paragraph_format.space_before = Pt(0)
        p_cert.paragraph_format.space_after = Pt(1.5)
        p_cert.add_run(" - ".join(selected_certs))

    # 8. Achievements (Smart selection of top 2-3 relevant)
    achievements = profile.get("achievements", [])
    if achievements:
        selected_achievements = select_relevant_items(achievements, jd_text, min_k=2, max_k=3)
        add_section_heading("Achievements & Leadership")
        p_ach = doc.add_paragraph()
        p_ach.paragraph_format.left_indent = Inches(0.1)
        p_ach.paragraph_format.space_before = Pt(0)
        p_ach.paragraph_format.space_after = Pt(1.5)
        p_ach.add_run(" - ".join(selected_achievements))

    # 9. Languages
    languages = profile.get("languages", {})
    if languages:
        add_section_heading("Languages")
        p_lang = doc.add_paragraph()
        p_lang.paragraph_format.left_indent = Inches(0.1)
        p_lang.paragraph_format.space_before = Pt(0)
        p_lang.paragraph_format.space_after = Pt(1)
        lang_items = [f"{lang} ({lvl})" for lang, lvl in languages.items()]
        p_lang.add_run(" - ".join(lang_items))

    # Save
    doc.save(output_path)


def create_cover_letter(job, profile, output_path):
    """Generates a customized cover letter preserving the exact base template, layout, and clickable links."""
    doc = docx.Document()
    
    # Page setup - Standard 0.75 inch clean margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # 1. Header (Left-aligned as in user's base template)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    run_name = title_p.add_run(profile["name"].upper())
    run_name.bold = True
    run_name.font.size = Pt(20)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(3)
    run_sub = sub_p.add_run("Full-Stack / Web Developer / Cloud Computing / AIML Intern")
    run_sub.bold = True
    run_sub.font.size = Pt(10.5)
    from docx.shared import RGBColor
    run_sub.font.color.rgb = RGBColor(29, 99, 184)

    contact_p1 = doc.add_paragraph()
    contact_p1.paragraph_format.space_after = Pt(1)
    run_c1 = contact_p1.add_run(f"{profile['phone']}  |  {profile['location']['current']}, India")
    run_c1.font.size = Pt(9.5)

    contact_p2 = doc.add_paragraph()
    contact_p2.paragraph_format.space_after = Pt(10)
    add_hyperlink(contact_p2, f"mailto:{profile['email']}", profile['email'])
    contact_p2.add_run("  |  ")
    add_hyperlink(contact_p2, f"https://{profile['linkedin']}", f"https://{profile['linkedin']}")
    contact_p2.add_run("  |  ")
    add_hyperlink(contact_p2, f"https://{profile['github']}", f"https://{profile['github']}")

    # Section Title: COVER LETTER
    cl_heading = doc.add_paragraph()
    cl_heading.paragraph_format.space_before = Pt(4)
    cl_heading.paragraph_format.space_after = Pt(6)
    run_cl = cl_heading.add_run("COVER LETTER")
    run_cl.bold = True
    run_cl.font.size = Pt(12)

    # Salutation
    p_salutation = doc.add_paragraph()
    p_salutation.paragraph_format.space_after = Pt(8)
    p_salutation.add_run("Dear Hiring Manager,")
    
    # Target tech stack derivation from JD
    jd_text = (job["title"] + " " + job.get("snippet", "")).lower()
    tech_phrase = "React and TypeScript"
    if "python" in jd_text or "django" in jd_text:
        tech_phrase = "Python, Django, and modern web frameworks"
    elif "node" in jd_text or "express" in jd_text:
        tech_phrase = "React, Node.js, and TypeScript"

    # Paragraph 1
    p1 = (
        f"I’m interested in applying for the {job['title']} at {job['company']}. "
        f"I’m especially drawn to the strong focus your team places on engineering fundamentals. "
        f"As an IT undergraduate, I’ve had the chance to build production-level, user-facing web applications using "
        f"{tech_phrase}. I’m now looking for an Opportunity where I can continue learning while also contributing in a meaningful way "
        f"to your Team or Company."
    )
    
    # Paragraph 2 (Project Spotlight)
    p2 = (
        "In my recent project, I’ve worked on a full-stack Web platform- ‘Counselling Centre, AVCOE’ in which I built a responsive front end "
        "using React and TypeScript. The goal was to make it easier for Students to access Support. I developed role-based dashboards for "
        "Students, Counsellor and Admin, handled state management, and worked with real-time data to make things smoother for both "
        "Students and Administrators. One feature I’m particularly proud of is a QR-based access system, which made it quicker for Students to "
        "get Help. I also added Reporting tools that simplified record-Handling process. Overall, this project helped me get better at building "
        "beautiful Interfaces, Managing States, and working closely with Back-end systems."
    )
    
    # Paragraph 3 (Culture, Growth & Japanese)
    p3 = (
        f"Beyond technical skills, I am highly motivated by {job['company']}’s culture of continuous growth and industrial exposure. Furthermore, I am "
        "currently learning Japanese (N5 level), which reflects my strong interest in cross-cultural communication and my long-term goal of "
        "building a career as an engineer in a global setting."
    )
    
    # Paragraph 4 (Closing)
    p4 = (
        "I’d welcome the opportunity to connect and contribute to your Team’s or Company’s goals. I appreciate you considering my "
        "application and hope to contribute and learn from your organization soon."
    )
    
    doc.add_paragraph(p1).paragraph_format.space_after = Pt(8)
    doc.add_paragraph(p2).paragraph_format.space_after = Pt(8)
    doc.add_paragraph(p3).paragraph_format.space_after = Pt(8)
    doc.add_paragraph(p4).paragraph_format.space_after = Pt(12)
    
    # Sign off
    p_signoff = doc.add_paragraph()
    p_signoff.add_run(f"Regards,\n{profile['name']}\n{profile['phone']}\n{profile['email']}")
    
    doc.save(output_path)


def generate_application_bundle(jobs, output_dir="dist"):
    """Generates resumes and cover letters for a list of jobs, zipping them by company."""
    profile = load_profile()
    
    if os.path.exists(output_dir):
        shutil.rmtree(output_dir)
    os.makedirs(output_dir)
    
    generated_files = []
    
    for job in jobs:
        company_clean = job["company"].replace(" ", "_").replace("/", "_")
        role_clean = job["title"].replace(" ", "_").replace("/", "_").replace("(", "").replace(")", "")
        job_dir = os.path.join(output_dir, f"{company_clean}_{job['location']}")
        os.makedirs(job_dir, exist_ok=True)
        
        name_clean = profile["name"].replace(" ", "_")
        resume_name = f"{name_clean}_Resume_{company_clean}_{role_clean}.docx"
        cl_name = f"{name_clean}_CoverLetter_{company_clean}.docx"
        
        resume_path = os.path.join(job_dir, resume_name)
        cl_path = os.path.join(job_dir, cl_name)
        
        create_resume(job, profile, resume_path)
        create_cover_letter(job, profile, cl_path)
        
        generated_files.append((job["company"], job_dir))
        print(f"Generated documents for {job['company']} in {job_dir}")
        
    # Zip the entire dist directory
    import datetime
    zip_name = f"{name_clean}_Applications_{datetime.date.today().strftime('%Y%m%d')}.zip"
    zip_path = os.path.join(os.path.dirname(output_dir), zip_name)
    
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(output_dir):
            for file in files:
                file_path = os.path.join(root, file)
                # Save relative path in zip
                rel_path = os.path.relpath(file_path, output_dir)
                zipf.write(file_path, rel_path)
                
    print(f"Bundled all materials into: {zip_path}")
    return zip_path

if __name__ == "__main__":
    from search_engine import run_job_search
    print("Testing document generation on first search result...")
    jobs = run_job_search()
    if jobs:
        zip_res = generate_application_bundle([jobs[0]], output_dir="temp_dist")
        print(f"Success! Zip generated: {zip_res}")
        if os.path.exists("temp_dist"):
            shutil.rmtree("temp_dist")
        if os.path.exists(zip_res):
            os.remove(zip_res)
